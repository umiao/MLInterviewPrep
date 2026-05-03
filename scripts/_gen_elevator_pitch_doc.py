"""One-shot generator for the elevator-pitch + 节奏 Word doc (T-throwaway).

Output: logs/elevator_pitch_summary.docx
Scope: prepared system-design problems, EXCLUDING the 4 eBay projects (per
user request 2026-04-29 17:35 Discord). Three tiers: Uber 2 (deep) /
Pinterest 7 (mid) / generic SD 19 (1-line table).

Compact layout: 0.5in margins, 10pt body / 9pt tables / 11-12pt headings,
single spacing, no emoji (project rule), bold for key terms inline.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _set_cell_shading(cell, hex_color: str) -> None:
    """Background-color a table cell."""
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_default_font(doc: Document, font_name: str = "Microsoft YaHei", size_pt: int = 10) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    rpr = style.element.get_or_add_rPr()
    rfont = rpr.find(qn("w:rFonts"))
    if rfont is None:
        from docx.oxml import OxmlElement

        rfont = OxmlElement("w:rFonts")
        rpr.append(rfont)
    rfont.set(qn("w:eastAsia"), font_name)
    rfont.set(qn("w:ascii"), font_name)
    rfont.set(qn("w:hAnsi"), font_name)


def _add_h(doc: Document, text: str, level: int = 1, color: tuple[int, int, int] = (0x1F, 0x4E, 0x79)) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6 if level >= 2 else 8)
    p.paragraph_format.space_after = Pt(2)
    sizes = {1: 14, 2: 12, 3: 11}
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 10))
    run.font.color.rgb = RGBColor(*color)


def _add_inline(doc: Document, parts: list[tuple[str, bool]], space_after: int = 2) -> None:
    """Add a paragraph with mixed bold / regular runs.

    Each part is (text, is_bold).
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for text, is_bold in parts:
        run = p.add_run(text)
        run.bold = is_bold
        run.font.size = Pt(10)


def _add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _add_kv(doc: Document, label: str, body: str) -> None:
    """Compact 'Label: body' line with bold label."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)


# ---------- Content ----------

NARRATION_PRINCIPLE = [
    ("开场 5s — 主导节奏", "我先用 1 句话定义这是什么问题, 然后给 30s 的高层架构. 你想 deep dive 哪一块, 随时打断我."),
    ("Frame 10s — 抓住 framing", "讲一句反 anti-pattern 的关键 framing (Uber Promo: 不是 redemption rate 而是 incremental profit; Restaurant Rec: 不是 query→result 而是 home feed 个性化模块)."),
    ("High-level 20s — 3-5 个组件", "客户端 → 路由 → 召回(多通道) → 粗排 → 精排 → 重排/diversity → 日志/反馈环. 每个组件给 1 个 industry jargon (H3, two-tower, MMoE, MMR)."),
    ("Deep dive cue 10s — 主动暴露三处可深挖", "我可以 deep dive (a) feature 系统怎么防 train-serve skew, (b) MMoE 多任务 head 怎么设计, (c) 召回怎么 merge 多路—— 你想哪一个?"),
    ("收尾 5s — Senior signal", "明确说 trade-off 和 fallback. 例: 'graceful degradation: 任一路挂了 fallback 到 popularity baseline, 不能整体不可用'."),
]

UBER_PITCHES = [
    {
        "id": "Uber Eats — 餐厅推荐 (Restaurant Recommendation)",
        "subtitle": "Home feed 个性化模块 (不是 search, 不是 reorder)",
        "30s": (
            "Uber Eats App 主页『为你推荐』个性化模块. 50M DAU, peak QPS 5-10K, "
            "外卖比一般电商 8x peakier (lunch/dinner 双峰), 候选 ~500-1000 缩到 30 张卡. "
            "三层架构: 召回 (H3 geo + two-tower + popular) → 粗排 GBDT → 精排 MMoE+DIN multi-task → MMR diversity rerank, "
            "feature snapshot at serving 防 train-serve skew, graceful degradation 保 99.95% 可用."
        ),
        "60s": (
            "我把这道题 frame 成 home feed 主页『为你推荐』个性化模块, 区别于 search 和『再来一单』非个性化模块. "
            "Non-functional: P99 < 200ms, 99.95% 可用. 规模上 50M DAU、peak 5-10K QPS, "
            "关键洞察是外卖比普通电商 8x peakier (lunch + dinner 双峰), capacity 要按 8x 预留. "
            "架构按 candidate funnel 走: 多通道召回 (H3 hexagon geo index + two-tower ANN + inverted + popular fallback) "
            "→ hard filter 餐厅是否 open → 粗排轻量 GBDT 缩到 100 → 精排 MMoE+DIN multi-task (CTR/CVR/dwell 多 head) 缩到 30 "
            "→ rerank MMR diversity + business rules + extreme filter (ETA 太长直接砍). "
            "三层时间 budget: offline 天级训练 / nearline 秒级特征更新 / online ms 级 serving. "
            "Senior 信号: feature snapshot at serving 防 train-serve skew, graceful degradation timeout+fallback, "
            "Model 学连续信号 + Policy 兜极端 (两层不是二选一), 评估用 cluster A/B 解 spillover."
        ),
        "keywords": "H3 (Uber 自研 hexagon geo) · two-tower · MMoE (多任务 multi-gate) · DIN (Deep Interest Net) · MMR diversity · Michelangelo (Uber ML platform) · feature snapshot · graceful degradation · position bias · cluster A/B",
        "rhythm": [
            ("[0-5s] Frame", "主页 home feed 个性化模块, 不是 search."),
            ("[5-15s] 规模 + peakiness", "50M DAU, 5-10K peak QPS, 外卖 8x peakier — capacity 按 8x 预留."),
            ("[15-35s] Funnel 架构", "召回多通道 H3+two-tower+popular → hard filter open → GBDT 粗排 → MMoE+DIN 精排 → MMR rerank."),
            ("[35-45s] 三层防御", "Offline 训练 / Nearline 特征 / Online serving 三层时间预算; feature snapshot 防 skew; graceful degradation."),
            ("[45-55s] Senior signal", "Model 学连续信号 + Policy 兜极端; cluster A/B 解 spillover; logging snapshot 是核心组件不是 nice-to-have."),
            ("[55-60s] Deep dive 主动暴露", "可以深挖: (a) MMoE expert 共享 + gate 设计 / (b) feature snapshot 实现 / (c) 多通道召回 merge 策略."),
        ],
    },
    {
        "id": "Uber — Budget-Constrained Promo Recommendation",
        "subtitle": "uplift modeling × constrained optimization 复合问题",
        "30s": (
            "给定固定 promo budget (例: $10M/月), 决定给哪些用户发什么 promo 最大化 incremental profit. "
            "ML 层用 uplift model (T/X-learner + XGBoost) 学 individual treatment effect τ(u,p); "
            "优化层 formulate 成 MCKP, 用 Lagrangian relaxation 在 N=10M 用户 scale 下并行求解 — "
            "每个用户独立 argmax τ - λc, 外层 binary search shadow price λ; "
            "评估用 long-running holdout + IPS/DR off-policy estimator."
        ),
        "60s": (
            "Frame 关键: 这不是 'predict redemption rate', 是 'maximize incremental profit'. "
            "踩坑点 = incrementality trap: 直接预测 P(redeem|u,p) 会偏向本来就活跃的用户, 等于 cannibalize 自有 GMV. "
            "正确 framing 是因果 treatment effect τ(u,p) = E[Y|do(T=p),X] - E[Y|do(T=0),X] (Pearl do-operator). "
            "三层架构: "
            "(1) ML 层 — randomized experiment 数据训 uplift model, T-learner / X-learner (Künzel 2019) + XGBoost, K 个 promo 选 multi-T 还是 S-learner 看 K 大小; "
            "(2) 优化层 — 单期 budget B 下 formulate Multiple-Choice Knapsack ILP, scale 上用 Lagrangian relaxation decouple 到 user 级独立 argmax τ(u,p)-λc(u,p), 外层对 shadow price λ 做 binary search 收敛到 budget 等式; "
            "PID controller 做 budget pacing 防止前几天烧光. "
            "(3) 探索层 — contextual bandit / Thompson sampling 在 uplift 输出上加 explore-exploit. "
            "评估: offline 用 IPS / Doubly Robust estimator 做 off-policy evaluation, "
            "online 用 long-running holdout (永远不发 promo 的小流量 control) 量真实 incremental profit. "
            "Senior 信号: 一上来就 frame 成 incrementality + causal, 写出 ILP formulation, 提 Lagrangian 强调 decouple, 提 OPE."
        ),
        "keywords": (
            "uplift / ITE / CATE · T-learner / X-learner / S-learner · MCKP (Multiple-Choice Knapsack) · "
            "Lagrangian relaxation · shadow price (λ) · PID budget pacing · contextual bandit · "
            "Thompson sampling · IPS (Inverse Propensity Scoring) · DR (Doubly Robust) · OPE (Off-Policy Evaluation) · "
            "long-running holdout · cannibalization · do-operator (Pearl)"
        ),
        "rhythm": [
            ("[0-5s] Frame", "最大化 incremental profit, 不是 redemption rate."),
            ("[5-15s] Anti-pattern 警告", "直接预测 redemption 等于 cannibalize 活跃用户的自有 GMV — 必须 frame 成 causal treatment effect."),
            ("[15-30s] 三层架构 ML/优化/探索", "ML: T/X-learner + XGBoost 学 τ(u,p); Opt: MCKP ILP + Lagrangian decouple 到独立 user; Explore: contextual bandit."),
            ("[30-40s] Scale 解法", "Lagrangian relaxation 关键: τ - λc 让每个用户独立 argmax, λ 外层 binary search → 10M 用户并行可解."),
            ("[40-50s] 评估", "Offline IPS/DR off-policy + Online long-running holdout 量真 incremental, A/B 不行因为 spillover."),
            ("[50-60s] Senior signal + Deep dive cue", "提 PID budget pacing, do-operator framing; 可深挖 (a) X-learner vs T-learner 选型 / (b) DR estimator 数学 / (c) bandit cold-start."),
        ],
    },
]

PINTEREST_PITCHES = [
    {
        "id": "Pinterest — Ad CTR Prediction",
        "30s": (
            "Pinterest promoted-pin CTR. 关键 clarify: ad 形态 (promoted pin / shopping / video / carousel)、surface (home/search/related)、"
            "目标是 ranking 还是 oCPM pricing (后者必须 calibrated). "
            "架构: feature store (user/pin/context/cross) → DCN-V2 / DeepFM / MaskNet → "
            "calibration (isotonic / Platt) → exploration (Thompson). 评估: AUC + Log-loss + Calibration error + online A/B revenue lift."
        ),
        "keywords": "DCN-V2 · DeepFM · MaskNet · feature crossing · calibration (Platt / isotonic) · oCPM · pacing · counterfactual logging",
        "rhythm": "Clarify 60s → high-level 30s → deep dive feature crossing 或 calibration → 提 counterfactual logging 防 selection bias.",
    },
    {
        "id": "Pinterest — User & Item (Pin) Embeddings",
        "30s": (
            "多消费场景: candidate gen ANN + ranking feature + similar-pins. 多目标 multi-task (repin/click/long-click/hide). "
            "架构: GraphSAGE 在 user-pin-board 异构图上学 pin embedding (Pinterest 经典 PinSage); "
            "user side 用 two-tower + 行为序列 attention. 冷启动用 content tower (CLIP image + text BERT) 兜底. "
            "训练: in-batch sampled softmax + hard negative mining. Serving: ANN (HNSW / IVF-PQ) 离线灌库."
        ),
        "keywords": "PinSage · GraphSAGE · two-tower · in-batch negatives · hard negative mining · HNSW · IVF-PQ · CLIP cold-start · multi-task",
        "rhythm": "Clarify 消费场景 → 单 vs 多 entity → 协同 vs 语义 → cold-start 兜底.",
    },
    {
        "id": "Pinterest — Personalized Chat Bot Recommending Pins",
        "30s": (
            "LLM chatbot 推 pins. 关键: LLM 不做 retrieval (避免幻觉 + 实时库存盲), 做 query understanding + 意图 artifact 生成, "
            "下游走传统 ANN 检索 (PinSage embedding + HNSW), 再用 LLM 做 grounded explanation. "
            "RAG 模式: pin metadata + user history 进 prompt context. 多模态: 图片 upload 走 CLIP encoder. "
            "Latency: streaming first-token < 500ms, 完整回答 < 3s."
        ),
        "keywords": "RAG · grounded generation · CLIP multimodal · function calling · streaming first-token · prompt cache · hallucination guardrails",
        "rhythm": "Frame LLM as artifact generator (借鉴 eBay LLM Orch 经验) → retrieval 走传统栈 → grounded reasoning → multimodal optional.",
    },
    {
        "id": "Pinterest — Pin Ranking for Home/Topic Feed",
        "30s": (
            "500M MAU, peak 100K QPS, P99 < 400ms (ranking budget ~150ms). Multi-objective: session length + repin + long-term retention. "
            "架构: 多通道召回 (PinSage ANN + 热门 + 关注 board + ML2) → MMoE 精排多 head (CTR/repin/long-click/hide) → "
            "diversity rerank (DPP 或 MMR) + frequency cap. Senior 信号: position bias debias (PAL / shallow tower), "
            "user behavior sequence model (DIN-like attention), Pareto multi-objective 加权."
        ),
        "keywords": "PinSage ANN · MMoE multi-head · DPP diversity · MMR · DIN attention · PAL position debias · Pareto weighting · long-term retention reward",
        "rhythm": "Clarify surface + scale → funnel 召回→粗排→精排→rerank → 多目标权重 trade-off → debias 收尾.",
    },
    {
        "id": "Pinterest — Pins Search Engine",
        "30s": (
            "Keyword + visual search (Lens). 数十 B pins, 100K QPS, P99 < 500ms. "
            "架构: query understanding (intent classify + spell correct + expansion) → 多通道召回 (inverted index BM25 + dense ANN HNSW + visual ANN) "
            "→ 粗排 lite tower → 精排 cross-attention BERT-style → diversity rerank. "
            "Lens 走 CLIP image embedding 同 ANN 索引. Personalization: user embedding 进 ranking feature 而非召回."
        ),
        "keywords": "BM25 · dense ANN HNSW · CLIP visual search · query expansion · cross-attention reranker · BERT-style · Lens",
        "rhythm": "Clarify text vs image → query understanding → 多通道召回 → 精排 → personalization 在 ranking 层不在召回.",
    },
    {
        "id": "Pinterest — Notification Recommendation",
        "30s": (
            "数十亿日发送, 多 channel (push/email/in-app). 核心: 不是 'how to send' 是 'who to send + what + when + how often'. "
            "Frame 成 utility maximization 减去 negative externality (annoyance / unsubscribe risk). "
            "架构: candidate generator (新 pin / followed board update / re-engagement) → relevance scorer (uplift on engagement) → "
            "frequency cap + send-time optimization (per-user circadian model) → multi-armed bandit on channel. "
            "Eval: online metric 是 long-term DAU lift 不是单封 CTR."
        ),
        "keywords": "uplift on engagement · frequency cap · send-time optimization · multi-armed bandit on channel · annoyance cost · long-term DAU lift",
        "rhythm": "Frame utility - annoyance → who/what/when/how-often 四问 → uplift 而非 CTR → 长期 metric.",
    },
    {
        "id": "Pinterest — Catalog Bulk Update (500M records, S3 + Async Fan-out)",
        "30s": (
            "数据平台问题不是 CRUD. 500M 记录 daily delta, 多下游 (search index / ranking feature store / ad serving). "
            "架构: S3 drop 触发 → Spark job 做 schema validate + dedup + diff (本期 vs 上期) → "
            "Kafka topic 按 entity 类型分 partition fan-out → 各下游消费幂等 upsert. "
            "Senior 信号: idempotency key (record_id + version), 增量 vs 全量 trade-off, exactly-once 用 transactional Kafka + 下游 dedup table."
        ),
        "keywords": "S3 drop · Spark validate/dedup/diff · Kafka fan-out · idempotency key · transactional producer · exactly-once · backpressure",
        "rhythm": "Clarify scale + freshness + 一致性 → batch vs streaming → fan-out partition 策略 → idempotency 兜底.",
    },
]

GENERIC_19 = [
    ("URL Shortener", "Hash/encode 长 URL + redirect. base62 encode of 64-bit ID, collision via retry. 读重 → cache + CDN, write 用 ID generator (Snowflake)."),
    ("Rate Limiter", "Token bucket / sliding window log / sliding window counter. 分布式用 Redis Lua atomic, 多区域 eventually consistent. Trade-off: 精度 vs latency."),
    ("Distributed Cache", "Consistent hashing 虚拟节点 + LRU/LFU/TinyLFU eviction; cache-aside vs write-through; singleflight 防 stampede; Bloom filter 防 penetration."),
    ("Notification System", "多 channel (push/SMS/email) priority queue, per-user rate limit, template render, retry 指数退避 + DLQ; provider abstraction."),
    ("News Feed (Instagram)", "Fan-out on write (普通用户) vs fan-out on read (celebrity 名人) hybrid. EdgeRank-style ML ranking. cursor pagination."),
    ("Chat (WhatsApp)", "WebSocket 持久连接, message ID Snowflake 单调有序, group chat fan-out, presence via heartbeat + Redis pub/sub."),
    ("Live Comments (Facebook)", "Fan-out tree (百万并发观看), SSE/WebSocket, 评论 sampling 防屏幕被刷屏, pre-moderation + reaction aggregation."),
    ("Real-time Game Leaderboard", "Redis Sorted Set ZADD/ZREVRANK, score-range sharding for 50M+ players, daily/weekly/season boards 分 key, Kafka peak shaving."),
    ("Ride-sharing (Uber)", "H3 geo index (Uber 自研), real-time driver location 写 Redis geospatial, 匹配走 Hungarian / greedy + ETA model, surge pricing 走 supply/demand 微观经济模型."),
    ("Proximity Service (Yelp)", "Geohash vs QuadTree vs R-Tree trade-off, radius query, 99:1 read-heavy → 多级 cache (CDN + Redis + L1)."),
    ("Search Autocomplete", "Trie + top-K 频率, real-time update 走 Kafka, prefix cache 在 client + edge; 多语言 fallback."),
    ("Top-K Heavy Hitters", "Count-Min Sketch + Min-Heap streaming, 三层聚合 (local → partition → global), Lambda Architecture 配合 hourly batch calibrate."),
    ("Ad Click Aggregator", "Lambda Architecture: Kafka + Flink exactly-once, 两层 dedup (Bloom + RocksDB), real-time fraud detection, ClickHouse OLAP, 批 billing 对账."),
    ("Web Crawler", "URL frontier, politeness (per-domain rate limit), dedup (URL + content fingerprint), 分布式协调 ZooKeeper, 重试 + 失败隔离."),
    ("Video Streaming (YouTube)", "上传 + 转码 DAG 并行 GPU, ABR adaptive bitrate (HLS/DASH), 三层 CDN (Edge POP → Shield → Origin S3), multi-CDN failover."),
    ("Cloud Storage (Dropbox)", "Block-level chunking (CDC / Rabin Fingerprint) + dedup + delta sync, conflict resolution (version vector), metadata DB sharded MySQL, WebSocket sync."),
    ("Price Drop Tracker", "Scraping pipeline (proxy rotation + golden test 防反爬), TimescaleDB 价格历史 (downsample + continuous aggregate), Z-score anomaly detection."),
    ("Online Judge (LeetCode)", "Sandbox (gVisor / Docker + cgroups + seccomp), submission queue RabbitMQ, test runner 早停, MOSS plagiarism (Winnowing), Redis Sorted Set 比赛 leaderboard."),
    ("Ticket Reservation", "SELECT FOR UPDATE SKIP LOCKED 座位锁, payment hold TTL, virtual queue 防秒杀挤爆, overbooking probability model, idempotent payment."),
]


def build_doc() -> Document:
    doc = Document()
    _set_default_font(doc)

    # narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("System Design Elevator Pitch & 讲述节奏 汇总")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(6)
    sub = p.add_run("覆盖范围: Uber 2 题 (深) · Pinterest 7 题 (中) · 通用 19 题 (速查)  ·  生成日期 2026-04-29")
    sub.italic = True
    sub.font.size = Pt(9)
    sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ---- Section 0: 节奏原则 ----
    _add_h(doc, "0. 通用讲述节奏 (60-90s 模板, 任何 SD 题都套这个)", level=1)
    for label, body in NARRATION_PRINCIPLE:
        _add_kv(doc, f"{label}  ", body)

    _add_inline(
        doc,
        [
            ("反 anti-pattern: ", True),
            ("(1) 上来就画图不澄清 → 死. (2) 不主动暴露 deep-dive cue → 面试官只能引导. (3) 不提 fallback / trade-off → 像中级. (4) 名词不展开 (说 'two-tower' 不解释) → 像背书.", False),
        ],
        space_after=8,
    )

    # ---- Section 1: Uber 2 ----
    _add_h(doc, "1. Uber 2 题 — Staff-level (深度 pitch)", level=1)
    for item in UBER_PITCHES:
        _add_h(doc, item["id"], level=2)
        _add_inline(doc, [(f"({item['subtitle']})", False)], space_after=2)

        _add_h(doc, "30s 版", level=3, color=(0x44, 0x6B, 0x99))
        _add_quote(doc, item["30s"])

        _add_h(doc, "60s 版", level=3, color=(0x44, 0x6B, 0x99))
        _add_quote(doc, item["60s"])

        _add_h(doc, "Keywords (展开名词解释, 面试时先说 name 再用一句话解释)", level=3, color=(0x44, 0x6B, 0x99))
        _add_inline(doc, [(item["keywords"], False)], space_after=4)

        _add_h(doc, "讲述节奏 (timestamp anchors)", level=3, color=(0x44, 0x6B, 0x99))
        for ts, body in item["rhythm"]:
            _add_kv(doc, f"{ts}  ", body)

    # ---- Section 2: Pinterest 7 ----
    _add_h(doc, "2. Pinterest 7 题 — 中等深度 pitch", level=1)
    for item in PINTEREST_PITCHES:
        _add_h(doc, item["id"], level=2)
        _add_quote(doc, item["30s"])
        _add_kv(doc, "Keywords  ", item["keywords"])
        _add_kv(doc, "节奏  ", item["rhythm"])
        # spacer
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(2)

    # ---- Section 3: Generic 19 速查表 ----
    _add_h(doc, "3. 通用 SD 19 题 — 一句话 pitch 速查", level=1)

    table = doc.add_table(rows=1 + len(GENERIC_19), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(5.2)

    hdr = table.rows[0].cells
    for i, h in enumerate(["题目", "30s 一句话 pitch (核心架构 + 关键 jargon)"]):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(hdr[i], "1F4E79")

    for r, (title, pitch) in enumerate(GENERIC_19, start=1):
        cells = table.rows[r].cells
        cells[0].text = ""
        cells[1].text = ""
        run0 = cells[0].paragraphs[0].add_run(title)
        run0.bold = True
        run0.font.size = Pt(9)
        run1 = cells[1].paragraphs[0].add_run(pitch)
        run1.font.size = Pt(9)
        if r % 2 == 0:
            _set_cell_shading(cells[0], "F2F2F2")
            _set_cell_shading(cells[1], "F2F2F2")

    # Footer note
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    f = p.add_run(
        "用法: 面试前 5 分钟扫一眼对应 section. 30s 版用于 'tell me about this problem'; "
        "60s 版用于深入 framing; keywords 是 'show, don't tell' 的弹药库 (说出来才算会). "
        "节奏 anchors 是 timestamps — 练几遍直到不看也能 pace."
    )
    f.italic = True
    f.font.size = Pt(9)
    f.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    return doc


def main() -> None:
    doc = build_doc()
    out_dir = Path("logs")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "elevator_pitch_summary.docx"
    doc.save(out_path)
    print(f"[OK] wrote {out_path.resolve()}")
    print(f"[OK] size = {out_path.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
